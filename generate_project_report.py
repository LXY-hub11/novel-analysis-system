"""
Generate comprehensive project report for the Novel Analysis System presentation.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import json


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2F5496')

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph()
    return table


def generate_report():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # ── Helper functions ──
    def add_heading(text, level=1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        return heading

    def add_para(text, bold=False, size=11, align=None, indent=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if align:
            p.alignment = align
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        return p

    def add_bullet(text, level=0):
        p = doc.add_paragraph(text, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return p

    def add_code_block(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        p.paragraph_format.left_indent = Cm(1)
        return p

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('多平台小说大数据分析\n与推荐系统')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('项目技术报告')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(47, 84, 150)

    for _ in range(4):
        doc.add_paragraph()

    info_lines = [
        f'编制日期：{datetime.now().strftime("%Y年%m月%d日")}',
        '版 本 号：V1.0',
        '技术栈：Django + Spark + Hadoop + ECharts',
        '文档性质：项目讲解报告',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.name = '微软雅黑'

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════
    add_heading('目  录', level=1)
    toc_items = [
        ('一、项目概述', ''),
        ('  1.1 项目背景与意义', ''),
        ('  1.2 项目目标', ''),
        ('  1.3 项目范围', ''),
        ('二、系统架构设计', ''),
        ('  2.1 整体架构', ''),
        ('  2.2 技术栈选型', ''),
        ('三、数据资源与采集', ''),
        ('  3.1 数据资源概况', ''),
        ('  3.2 多平台数据采集引擎', ''),
        ('  3.3 数据处理流程', ''),
        ('四、核心功能模块', ''),
        ('  4.1 Web应用功能概览', ''),
        ('  4.2 数据分析系统', ''),
        ('  4.3 推荐算法引擎', ''),
        ('五、数据可视化系统', ''),
        ('六、数据分析成果', ''),
        ('  6.1 平台对比分析', ''),
        ('  6.2 分类分布分析', ''),
        ('  6.3 数据洞察结论', ''),
        ('七、项目结构总览', ''),
        ('八、技术亮点与创新', ''),
        ('九、总结与展望', ''),
    ]
    for item, _ in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 1: Project Overview
    # ══════════════════════════════════════════════════════════════
    add_heading('一、项目概述', level=1)

    add_heading('1.1 项目背景与意义', level=2)
    add_para('中国网络文学市场经历了二十余年的高速发展，已成为全球最大的数字阅读市场之一。截至2024年底，中国网络文学用户规模突破5.8亿，市场规模超过400亿元人民币。市场上存在起点中文网、番茄小说、晋江文学城等十余个主流平台，各平台积累了海量的作品数据。', indent=True)
    add_para('然而，当前网络文学行业面临"数据孤岛"困境：各平台数据相互封闭，读者难以跨平台发现优质作品，作者难以全面把握市场趋势，出版机构和投资者也难以进行系统的竞品分析。本项目正是针对这一痛点，构建一个多平台小说数据聚合、分析与智能推荐的一站式平台。', indent=True)

    add_heading('1.2 项目目标', level=2)
    goals = [
        '实现8大主流网络文学平台的数据采集与整合，覆盖15+分类、600+作品数据',
        '建立基于Hadoop + Spark的大数据处理与分析框架，支持多维度数据挖掘',
        '构建基于Django的全栈Web应用，提供友好的中文用户交互界面',
        '实现基于TF-IDF内容分析 + 协同过滤 + 跨平台推荐的混合推荐算法',
        '提供10+种交互式数据可视化图表，直观展示多维度分析结果',
        '搭建完善的用户系统，支持收藏、评分、个性化推荐等功能',
    ]
    for g in goals:
        add_bullet(g)

    add_heading('1.3 项目范围', level=2)
    add_para('项目覆盖以下8个主流网络文学平台：', indent=True)
    platforms = [
        '起点中文网（阅文集团）—— 行业龙头，男频优势',
        '番茄小说（字节跳动）—— 免费阅读模式领导者',
        '晋江文学城 —— 女频内容王者',
        '纵横中文网 —— 老牌文学平台',
        '七猫小说 —— 免费阅读新势力',
        '飞卢小说 —— 同人/二次元特色',
        '17k小说网 —— 老牌原创文学平台',
        '潇湘书院 —— 女频文学重镇',
    ]
    for p in platforms:
        add_bullet(p)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 2: System Architecture
    # ══════════════════════════════════════════════════════════════
    add_heading('二、系统架构设计', level=1)

    add_heading('2.1 整体架构', level=2)
    add_para('系统采用经典的大数据Web应用六层架构设计，从上到下依次为：', indent=True)

    arch_data = [
        ('层级', '技术选型', '核心职责'),
        ('数据采集层', 'Python + Requests + BS4 + Scrapy', '多平台并发爬取，8平台数据统一采集'),
        ('数据存储层', 'Hadoop HDFS + MySQL 8.0 / SQLite', '分布式文件存储 + 关系型数据库'),
        ('数据处理层', 'Apache Spark 3.4 + Pandas + NumPy', '大数据清洗、统计分析、特征工程'),
        ('算法引擎层', 'Scikit-learn + 自研TF-IDF + 协同过滤', '内容分析、相似度计算、混合推荐'),
        ('Web应用层', 'Django 5.2 + MTV架构', '全栈Web框架、用户系统、业务逻辑'),
        ('可视化层', 'ECharts 5.0 + AJAX + jQuery', '10+种交互式图表、实时数据API'),
    ]
    add_styled_table(doc, arch_data[0], arch_data[1:])
    add_para('架构设计遵循"高内聚、低耦合"原则，各层之间通过标准化接口通信。数据采集层独立运行，通过JSON格式将数据传递给存储层；算法引擎层从数据库读取数据，计算结果通过API暴露给前端；可视化层通过AJAX异步获取数据，实现前后端分离。', indent=True)

    add_heading('2.2 技术栈选型', level=2)

    tech_data = [
        ('技术组件', '选型方案', '版本', '选型理由'),
        ('后端框架', 'Django', '5.2', '成熟稳定，ORM强大，中文支持好'),
        ('大数据计算', 'Apache Spark', '3.4.1', '业界标准，Python API友好'),
        ('分布式存储', 'Apache Hadoop', '3.3.6', 'Spark生态配套，HDFS存储'),
        ('关系数据库', 'MySQL / SQLite', '8.0 / 3', '双模式支持，灵活切换'),
        ('科学计算', 'NumPy + Pandas', '1.24+ / 2.0+', '数据处理和向量运算基础'),
        ('机器学习', 'Scikit-learn', '1.3+', '经典ML算法库'),
        ('数据可视化', 'ECharts', '5.0', '功能丰富，中文友好，交互性强'),
        ('前端框架', 'jQuery + Bootstrap', '3.x / 5.x', '轻量级，兼容性好'),
        ('爬虫框架', 'Requests + BS4', '2.31+ / 4.12+', '灵活可控，支持自定义解析'),
    ]
    add_styled_table(doc, tech_data[0], tech_data[1:])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 3: Data Resources
    # ══════════════════════════════════════════════════════════════
    add_heading('三、数据资源与采集', level=1)

    add_heading('3.1 数据资源概况', level=2)
    add_para('当前系统已积累丰富的小说数据资源，具体指标如下：', indent=True)

    data_summary = [
        ('数据指标', '数值', '说明'),
        ('小说总数', '600+', '覆盖8大平台完整数据'),
        ('覆盖平台', '8个', '主流网络文学平台全覆盖'),
        ('覆盖分类', '15个', '玄幻/仙侠/都市/科幻/历史/军事等'),
        ('独特作者', '74位', '跨平台去重统计'),
        ('总字数', '9.28亿+', '所有作品字数累计'),
        ('总点击量', '1,372万+', '跨平台点击数据汇总'),
        ('总收藏量', '181万+', '跨平台收藏数据汇总'),
        ('数据时间范围', '2026年2月-5月', '近4个月持续采集'),
        ('数据字段', '13个', '标题/作者/分类/简介/字数/状态/平台等'),
    ]
    add_styled_table(doc, data_summary[0], data_summary[1:])

    add_heading('3.2 多平台数据采集引擎', level=2)
    add_para('系统设计了一套统一的多平台数据采集框架，核心特点包括：', indent=True)
    features = [
        '配置化采集：每个平台的URL、CSS选择器、数据字段均通过PlatformConfig类配置管理，新增平台只需添加配置字典',
        '并发爬取：基于ThreadPoolExecutor实现多线程并发采集，显著提升数据采集效率',
        '反爬虫策略：内置随机User-Agent轮换、请求间隔随机化、自动重试机制',
        '数据清洗：采集后自动进行数据去重、空值填充、格式标准化处理',
        '分类覆盖：每个平台独立配置分类URL，确保各类别数据均衡采集',
    ]
    for f in features:
        add_bullet(f)
    add_para('每个平台支持采集的数据字段包括：标题、作者、分类、简介、字数、点击量、收藏数、推荐数、最新章节、更新状态等13个维度。', indent=True)

    add_heading('3.3 数据处理流程', level=2)
    add_para('数据从采集到入库的完整处理流水线：', indent=True)
    add_bullet('步骤1：爬虫引擎按平台/分类并发采集原始HTML页面')
    add_bullet('步骤2：BeautifulSoup解析HTML，提取结构化数据字段')
    add_bullet('步骤3：数据清洗与预处理（去重、空值处理、格式标准化）')
    add_bullet('步骤4：生成统一格式的JSON数据文件（novels_data.json）')
    add_bullet('步骤5：Spark读取JSON，进行大规模数据处理与特征工程')
    add_bullet('步骤6：通过import_to_db.py导入MySQL/SQLite数据库')
    add_bullet('步骤7：Django ORM提供数据访问接口，前端通过API获取')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 4: Core Features
    # ══════════════════════════════════════════════════════════════
    add_heading('四、核心功能模块', level=1)

    add_heading('4.1 Web应用功能概览', level=2)
    add_para('系统提供完整的Web应用功能，包含7大核心模块：', indent=True)

    modules = [
        ('功能模块', '页面路由', '核心功能描述'),
        ('首页', '/index/', '系统介绍、最新小说展示、热门排行、功能导航'),
        ('小说列表', '/novels/', '多条件筛选（分类/平台）、关键词搜索、分页浏览'),
        ('小说详情', '/novel/<id>/', '完整信息展示、收藏、五星评分、相似推荐'),
        ('分类分析', '/category-analysis/', '分类统计、饼图可视化、分类热度排行'),
        ('数据可视化', '/visualization/', '多维度图表、平台对比、分类分布、来源分析'),
        ('个性化推荐', '/recommendations/', '基于用户收藏的智能推荐、混合推荐策略'),
        ('用户中心', '/collections/', '注册登录、个人资料、收藏管理、活动历史'),
        ('增强分析', '/enhanced-analysis/', '高级分析仪表盘、跨平台深度对比'),
    ]
    add_styled_table(doc, modules[0], modules[1:])

    add_heading('4.2 数据分析系统', level=2)
    add_para('数据分析系统是项目的核心引擎，基于Apache Spark实现大规模数据处理：', indent=True)

    add_para('（1）平台对比分析', bold=True, size=10.5)
    add_para('对8大平台进行多维度对比，包括作品数量、总点击量、平均点击数、总收藏量、平均收藏量、总字数、最高点击量等指标，全面评估各平台的内容生态和用户活跃度。', size=10)

    add_para('（2）分类分布分析', bold=True, size=10.5)
    add_para('统计15个分类的小说数量分布、各类别平均点击量、平均收藏量，识别热门内容类型和蓝海分类机会。', size=10)

    add_para('（3）字数分布分析', bold=True, size=10.5)
    add_para('将作品按字数分为短篇(<20万)、中篇(20-50万)、长篇(50-100万)、超长篇(100-200万)、巨作(200-500万)、史诗(>500万)六个级别，分析各平台的内容长度策略。', size=10)

    add_para('（4）作者分析', bold=True, size=10.5)
    add_para('按作品数量、总点击量、平均点击量对Top 30作者进行排行，分析高产作者和人气作者的分布特征。', size=10)

    add_para('（5）更新状态分析', bold=True, size=10.5)
    add_para('统计各平台连载中与已完结作品的比例，分析完结率与平台策略的关系。', size=10)

    add_para('（6）热度相关性分析', bold=True, size=10.5)
    add_para('计算点击量与字数、点击量与收藏量、收藏量与推荐量之间的皮尔逊相关系数，挖掘影响作品热度的关键因素。', size=10)

    add_heading('4.3 推荐算法引擎', level=2)
    add_para('系统实现了融合多种策略的混合推荐算法：', indent=True)

    rec_data = [
        ('推荐策略', '权重', '算法说明'),
        ('协同过滤推荐', '35%', '基于用户收藏偏好，计算用户间与物品间相似度，挖掘潜在兴趣'),
        ('分类偏好推荐', '25%', '根据用户偏好的分类，在该分类下推荐高人气作品'),
        ('平台偏好推荐', '20%', '根据用户偏好的平台，结合该平台的流行度进行推荐'),
        ('全局热度推荐', '20%', '综合点击量、收藏量、收藏/点击转化率计算作品热度分'),
    ]
    add_styled_table(doc, rec_data[0], rec_data[1:])

    add_para('')
    add_para('算法核心创新点：', bold=True, size=11)
    highlights = [
        '自研TF-IDF向量化：针对中文小说简介的特点，使用字符级Bigram分词构建500维特征向量，实现基于内容的相似度计算（余弦相似度）',
        '跨平台推荐：支持用户指定目标平台，在目标平台中寻找与源平台小说内容相似的作品，打破平台壁垒',
        '多样性增强：通过同作者惩罚（diversity_factor=0.3）和同分类计数惩罚（0.5×factor）避免推荐结果同质化',
        '多属性融合：综合分类一致性(+3.0)、平台一致性(+1.5)、作者一致性(+2.5)、字数相似度(+1.0)四个维度计算物品相似度',
    ]
    for h in highlights:
        add_bullet(h)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 5: Data Visualization
    # ══════════════════════════════════════════════════════════════
    add_heading('五、数据可视化系统', level=1)
    add_para('系统基于ECharts 5.0构建了丰富的交互式数据可视化仪表盘，提供12个RESTful API端点支撑前端实时数据展示：', indent=True)

    api_data = [
        ('API端点', '请求方式', '功能说明'),
        ('/api/overview/', 'GET', '系统全局概览统计（总数/平台/分类/作者/状态）'),
        ('/api/platform-comparison/', 'GET', '8大平台多维度对比数据'),
        ('/api/category-stats/', 'GET', '分类统计含平台维度细分'),
        ('/api/category-platform/', 'GET', '分类×平台交叉分析'),
        ('/api/word-count-dist/', 'GET', '字数级别分布（整体+分平台）'),
        ('/api/popularity/', 'GET', '热门作品排行（支持多字段排序）'),
        ('/api/author-analysis/', 'GET', '作者产出与热度分析Top30'),
        ('/api/status-analysis/', 'GET', '更新状态分平台统计'),
        ('/api/scatter/', 'GET', '散点图数据（字数vs点击量，平台着色）'),
        ('/api/rating-analysis/', 'GET', '用户评分分析'),
        ('/api/search/', 'GET', '高级搜索（多条件+排序）'),
        ('/api/platforms-list/', 'GET', '筛选器下拉列表数据'),
    ]
    add_styled_table(doc, api_data[0], api_data[1:])
    add_para('前端通过AJAX异步调用上述API获取JSON数据，再由ECharts渲染为散点图、热力图、堆叠柱状图、饼图、玫瑰图、仪表盘等10+种图表类型，支持图表切换、数据缩放、tooltip详情等交互功能。', indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 6: Data Analysis Results
    # ══════════════════════════════════════════════════════════════
    add_heading('六、数据分析成果', level=1)

    add_heading('6.1 平台对比分析', level=2)
    add_para('基于600条真实数据的平台对比分析结果：', indent=True)

    plat_data = [
        ('平台', '作品数', '总点击量', '总收藏量', '平均点击量', '总字数(万)'),
        ('番茄小说', '74', '2,896,031', '365,553', '39,136', '8,917'),
        ('起点中文网', '85', '2,346,807', '336,073', '27,609', '11,879'),
        ('七猫小说', '69', '1,895,988', '245,244', '27,478', '10,415'),
        ('晋江文学城', '72', '1,530,834', '219,153', '21,262', '11,221'),
        ('飞卢小说', '70', '1,441,694', '193,991', '20,596', '10,785'),
        ('潇湘书院', '84', '1,222,057', '168,814', '14,548', '12,998'),
        ('纵横中文网', '71', '1,036,444', '147,162', '14,599', '9,325'),
        ('17k小说网', '75', '1,005,435', '133,706', '13,406', '13,053'),
    ]
    add_styled_table(doc, plat_data[0], plat_data[1:])
    add_para('关键发现：番茄小说虽然作品数不是最多，但总点击量(289万)和平均点击量(3.9万)均排名第一，体现了免费阅读模式的流量优势。起点中文网以85部作品数排名第一，反映其内容生态最为丰富。', indent=True)

    add_heading('6.2 分类分布分析', level=2)

    cat_data = [
        ('排名', '分类', '作品数量', '占比', '典型平台'),
        ('1', '都市', '116', '19.3%', '起点/番茄/七猫'),
        ('2', '玄幻', '70', '11.7%', '起点/纵横'),
        ('3', '言情', '55', '9.2%', '晋江/潇湘'),
        ('4', '仙侠', '50', '8.3%', '起点/17k'),
        ('5', '轻小说', '49', '8.2%', '飞卢/番茄'),
        ('6', '科幻', '44', '7.3%', '起点/七猫'),
        ('7', '奇幻', '43', '7.2%', '晋江/起点'),
        ('8', '历史', '42', '7.0%', '起点/纵横'),
    ]
    add_styled_table(doc, cat_data[0], cat_data[1:])
    add_para('都市类以116部作品占比19.3%遥遥领先，反映都市题材在当前网络文学市场中的绝对主流地位。玄幻和言情分列二三位，分别代表了男频和女频的核心内容方向。', indent=True)

    add_heading('6.3 数据洞察结论', level=2)
    insights = [
        '平台流量集中度：番茄小说和起点中文网合计贡献了总点击量的38%，呈现"双寡头"格局',
        '内容长度差异：17k小说网(1305万字)平均单作品字数最高(17.4万)，番茄小说(892万字)最低(12.0万)，反映平台内容策略差异',
        '完结率分析：440部连载中 vs 160部已完结，整体完结率约26.7%，"连载中"是网络文学的主要状态',
        '收藏转化率：番茄小说收藏/点击比(12.6%)最高，反映其推荐算法和用户体验的精准性',
        '跨平台差异：各平台在内容类型、作品长度、用户活跃度等维度存在显著差异，验证了跨平台分析的价值',
    ]
    for insight in insights:
        add_bullet(insight)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 7: Project Structure
    # ══════════════════════════════════════════════════════════════
    add_heading('七、项目结构总览', level=1)

    add_code_block('小说分析系统/')
    add_code_block('├── novels_data.json           # 600+小说原始数据(JSON)')
    add_code_block('├── generate_report.py         # 可行性报告生成脚本')
    add_code_block('├── 可行性分析报告.docx          # 已生成的可行性分析报告')
    add_code_block('├── hadoop-config.cmd          # Hadoop环境配置脚本')
    add_code_block('├── hadoop-env.cmd             # Hadoop环境变量脚本')
    add_code_block('│')
    add_code_block('└── novel_project/             # Django项目主目录')
    add_code_block('    ├── manage.py              # Django管理脚本')
    add_code_block('    ├── requirements.txt       # Python依赖(11个库)')
    add_code_block('    ├── db.sqlite3             # SQLite数据库')
    add_code_block('    ├── install.bat            # Windows一键安装脚本')
    add_code_block('    ├── start_server.bat        # 快速启动脚本')
    add_code_block('    ├── check_env.py            # 环境检测脚本')
    add_code_block('    ├── init_project.py         # 项目初始化脚本')
    add_code_block('    │')
    add_code_block('    ├── novel_analysis/         # Django配置目录')
    add_code_block('    │   ├── settings.py         # 项目配置(SQLite/MySQL双模式)')
    add_code_block('    │   ├── urls.py             # 根URL路由')
    add_code_block('    │   ├── wsgi.py             # WSGI部署配置')
    add_code_block('    │   └── asgi.py             # ASGI异步配置')
    add_code_block('    │')
    add_code_block('    ├── apps/                   # Django应用目录')
    add_code_block('    │   ├── novels/             # 小说管理应用(核心)')
    add_code_block('    │   │   ├── models.py       # 6个数据模型')
    add_code_block('    │   │   ├── views.py        # 15个视图+12个API端点')
    add_code_block('    │   │   ├── urls.py         # URL路由配置')
    add_code_block('    │   │   ├── admin.py        # Django管理后台')
    add_code_block('    │   │   ├── templates/      # 7个HTML模板页面')
    add_code_block('    │   │   └── static/         # CSS + JS + ECharts')
    add_code_block('    │   └── users/             # 用户管理应用')
    add_code_block('    │       ├── models.py       # 用户模型')
    add_code_block('    │       ├── views.py        # 注册/登录/资料')
    add_code_block('    │       └── templates/      # 用户相关页面')
    add_code_block('    │')
    add_code_block('    ├── spark_scripts/          # Spark大数据处理脚本')
    add_code_block('    │   ├── spark_config.py     # Spark配置(4G内存/Kryo序列化)')
    add_code_block('    │   ├── multi_platform_crawler.py  # 8平台数据采集引擎')
    add_code_block('    │   ├── comprehensive_analysis.py  # 综合数据分析(6维分析)')
    add_code_block('    │   ├── enhanced_recommendation.py # 增强推荐算法(TF-IDF+混合)')
    add_code_block('    │   ├── generate_enriched_data.py   # 数据增强生成')
    add_code_block('    │   ├── recommendation.py   # 基础推荐算法')
    add_code_block('    │   ├── import_to_db.py     # 数据导入数据库')
    add_code_block('    │   └── fix_data.py         # 数据修复工具')
    add_code_block('    │')
    add_code_block('    ├── data/                   # 分析输出数据')
    add_code_block('    │   ├── enriched_novels.json   # 增强后的小说数据')
    add_code_block('    │   └── analysis_report.json   # Spark分析报告输出')
    add_code_block('    │')
    add_code_block('    └── logs/                   # 系统日志')
    add_code_block('        └── django.log')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 8: Technical Highlights
    # ══════════════════════════════════════════════════════════════
    add_heading('八、技术亮点与创新', level=1)

    highlights_list = [
        ('1. 配置化的多平台采集框架', '每个平台独立配置，新增平台仅需添加约30行配置代码（PlatformConfig字典），无需修改核心采集逻辑。支持CSS选择器配置化、并发采集、自动重试、反爬虫策略，实现了8大平台的统一数据采集。'),
        ('2. 自研中文TF-IDF向量化引擎', '针对中文小说简介的文本特点，设计字符级Unigram+Bigram混合分词策略，无需依赖jieba等外部分词库，提取500维特征向量。使用余弦相似度计算作品间的内容相似性，为跨平台推荐提供基础。'),
        ('3. 四策略融合混合推荐算法', '将协同过滤(35%)、分类偏好(25%)、平台偏好(20%)、全局热度(20%)四种策略加权融合，并通过多样性增强机制（同作者惩罚、同分类计数惩罚）避免推荐结果的"信息茧房"效应。首次实现了跨平台的内容推荐能力。'),
        ('4. Spark + Python双模式分析引擎', '核心分析逻辑在Spark上实现，同时提供Python原生回退方案（使用defaultdict和Counter实现相同功能）。当Spark环境不可用时自动切换到Python模式，保证系统的鲁棒性和可移植性。'),
        ('5. SQLite/MySQL双数据库模式', '开发阶段使用SQLite（零配置），生产环境可无缝切换至MySQL 8.0，仅需取消settings.py中的注释即可。兼顾开发便捷性和生产可靠性。'),
        ('6. 前后端分离的API架构', '前端HTML页面与后端通过12个RESTful API端点进行数据交互，所有图表数据通过AJAX异步加载。API支持丰富的查询参数（排序、筛选、分页），为后续开发移动端或第三方接入预留接口。'),
        ('7. 丰富的可视化图表体系', '基于ECharts 5.0实现散点图（字数vs热度）、堆叠柱状图（平台×分类）、饼图（分类分布）、玫瑰图（平台占比）、热力图（多维交叉分析）等10+种图表类型，支持自适应布局和动态筛选。'),
    ]
    for title_text, desc in highlights_list:
        add_para(title_text, bold=True, size=11)
        add_para(desc, size=10.5, indent=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 9: Conclusion
    # ══════════════════════════════════════════════════════════════
    add_heading('九、总结与展望', level=1)

    add_heading('9.1 项目成果总结', level=2)
    add_para('本项目成功构建了一个完整的多平台小说大数据分析与推荐系统，取得以下核心成果：', indent=True)

    results = [
        '数据层面：完成8大平台、15个分类、600+部小说的数据采集与整合，总数据覆盖928万+字数、1372万+点击量、181万+收藏量',
        '技术层面：搭建了完整的Spark + Hadoop + Django + ECharts技术栈，实现了从数据采集→处理→分析→展示的完整流水线',
        '算法层面：自研TF-IDF向量化引擎+四策略融合混合推荐算法，支持跨平台内容发现和多样性推荐',
        '应用层面：构建了包含首页、列表、详情、分析、可视化、推荐、用户中心7大模块的完整Web应用',
        'API层面：提供12个RESTful数据接口，覆盖平台对比、分类分析、作者分析、热度排行等多维度数据',
        '文档层面：产出环境搭建指南(PROJECT_GUIDE.md)、可行性分析报告、技术报告等完整文档',
    ]
    for r in results:
        add_bullet(r)

    add_heading('9.2 技术指标汇总', level=2)

    metrics = [
        ('类别', '指标', '数值'),
        ('数据规模', '小说总数', '600+'),
        ('数据规模', '覆盖平台数', '8个'),
        ('数据规模', '覆盖分类数', '15个'),
        ('数据规模', '数据字段数', '13个'),
        ('系统功能', 'Web页面数', '8个'),
        ('系统功能', 'RESTful API端点', '12个'),
        ('系统功能', '数据模型数', '6个'),
        ('系统功能', '可视化图表类型', '10+种'),
        ('算法能力', '推荐策略数', '4种'),
        ('算法能力', 'TF-IDF特征维度', '500维'),
        ('算法能力', '支持推荐类型', '5种(内容/协同/跨平台/热度/混合)'),
        ('代码规模', 'Python脚本文件', '15+个'),
        ('代码规模', 'HTML模板文件', '10+个'),
    ]
    add_styled_table(doc, metrics[0], metrics[1:])

    add_heading('9.3 未来展望', level=2)
    future = [
        '引入大语言模型（LLM）：利用LLM进行深度的内容理解、小说摘要生成、写作风格分析和智能书评生成，提升内容分析的深度和精度',
        '实时数据管道：构建基于Kafka/Flume的实时数据采集管道，实现新作品上架、数据更新的准实时监控',
        '用户画像深化：引入用户行为追踪，构建细粒度用户画像，实现更深度的个性化推荐',
        '社区功能扩展：增加书评、书单、讨论组等社交功能，从数据分析平台升级为读者社区平台',
        '商业化探索：逐步推出VIP订阅、API服务、定制报告、企业合作等商业化产品',
        '移动端适配：开发React Native或微信小程序版本，触达更广泛的移动端用户群体',
        'AIGC辅助创作：利用生成式AI技术为作者提供创作灵感、大纲优化、读者偏好分析等辅助工具',
    ]
    for f in future:
        add_bullet(f)

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 报告结束 —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # ── Save ──
    output_path = r'D:\项目\小说分析系统\项目技术报告.docx'
    doc.save(output_path)
    print(f'项目技术报告已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    generate_report()
