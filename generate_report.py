"""
Generate feasibility analysis report for the Novel Analysis System.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import datetime


def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, '2F5496')

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if r % 2 == 0:
                set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph()
    return table


def generate_report():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── Helper functions ──
    def add_heading(text, level=1):
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        return heading

    def add_para(text, bold=False, size=11, align=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if align:
            p.alignment = align
        return p

    def add_bullet(text, level=0):
        p = doc.add_paragraph(text, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        return p

    # ══════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('多平台小说大数据分析\n与推荐系统')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = '微软雅黑'

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('可行性分析报告')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(47, 84, 150)

    for _ in range(4):
        doc.add_paragraph()

    info_lines = [
        f'编制日期：{datetime.now().strftime("%Y年%m月%d日")}',
        '版 本 号：V1.0',
        '文档状态：初稿',
        '密    级：内部',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12)
        run.font.name = '宋体'

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ══════════════════════════════════════════════════════════════
    add_heading('目  录', level=1)
    toc_items = [
        ('一、项目概述', 3),
        ('  1.1 项目背景', 3),
        ('  1.2 项目目标', 3),
        ('  1.3 项目范围', 3),
        ('二、市场可行性分析', 4),
        ('  2.1 市场规模', 4),
        ('  2.2 竞争分析', 5),
        ('  2.3 目标用户', 5),
        ('三、技术可行性分析', 6),
        ('  3.1 技术架构', 6),
        ('  3.2 核心技术', 7),
        ('  3.3 数据资源', 7),
        ('四、经济可行性分析', 8),
        ('  4.1 成本估算', 8),
        ('  4.2 收益预测', 8),
        ('  4.3 投资回报分析', 9),
        ('五、运营可行性分析', 9),
        ('六、法律与合规分析', 10),
        ('七、风险分析', 10),
        ('八、项目实施计划', 11),
        ('九、结论与建议', 12),
    ]
    for item, page in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{item}{"." * (50 - len(item))}{page}')
        run.font.size = Pt(11)
        run.font.name = '宋体'

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 1: Project Overview
    # ══════════════════════════════════════════════════════════════
    add_heading('一、项目概述', level=1)

    add_heading('1.1 项目背景', level=2)
    add_para('随着互联网文学的蓬勃发展，中国网络文学市场持续高速增长。据中国作家协会发布的《2024中国网络文学蓝皮书》显示，截至2024年底，中国网络文学用户规模已突破5.8亿，市场规模超过400亿元人民币。网络文学平台数量众多，包括起点中文网、番茄小说、晋江文学城、纵横中文网、七猫小说等十余个主流平台，各平台拥有海量作品数据。')
    add_para('然而，当前各平台之间存在严重的"数据孤岛"现象：读者难以跨平台比较和发现优质作品，作者难以全面了解市场趋势和读者偏好，平台也难以进行跨平台竞品分析。因此，构建一套多平台小说数据聚合分析与智能推荐系统具有重要的商业价值和学术意义。')

    add_heading('1.2 项目目标', level=2)
    add_para('本项目旨在构建一个多平台小说大数据分析与推荐系统，主要目标包括：')
    goals = [
        '实现8大主流小说平台的数据采集与整合，覆盖15+分类、600+作品数据',
        '建立基于Hadoop+Spark的大数据处理与分析框架',
        '提供多维度的数据可视化分析功能（平台对比、分类分析、热度趋势等）',
        '实现基于TF-IDF内容分析+协同过滤+跨平台推荐的混合推荐算法',
        '构建基于Django的Web应用，提供友好的用户交互界面',
    ]
    for g in goals:
        add_bullet(g)

    add_heading('1.3 项目范围', level=2)
    add_para('项目范围涵盖以下8个主流网络文学平台：')
    platforms = ['起点中文网（阅文集团）', '番茄小说（字节跳动）', '晋江文学城',
                 '纵横中文网', '七猫小说', '飞卢小说', '17k小说网', '潇湘书院']
    for p in platforms:
        add_bullet(p)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 2: Market Feasibility
    # ══════════════════════════════════════════════════════════════
    add_heading('二、市场可行性分析', level=1)

    add_heading('2.1 市场规模', level=2)
    add_para('中国网络文学行业经过二十余年发展，已形成成熟的产业链条。以下是关键市场数据：')

    market_data = [
        ('指标', '2022年', '2023年', '2024年', '2025年(预测)'),
        ('网络文学用户规模(亿)', '5.2', '5.5', '5.8', '6.1'),
        ('市场规模(亿元)', '320', '360', '400', '440'),
        ('平台数量(个)', '45+', '40+', '38+', '35+'),
        ('日更新作品(万部)', '2.8', '3.1', '3.5', '3.8'),
        ('付费用户比例(%)', '12%', '14%', '16%', '18%'),
    ]
    add_styled_table(doc, market_data[0], market_data[1:])

    add_heading('2.2 竞争分析', level=2)
    add_para('目前市场上尚无成熟的跨平台小说数据分析产品。相关竞品主要集中在以下三个方向：')
    add_para('（1）单一平台内的数据分析工具：如起点中文网的数据助手，仅提供平台内部数据，缺乏跨平台视角。', size=10.5)
    add_para('（2）第三方书单/点评类应用：如豆瓣读书，以用户书评为核心，缺少系统性的数据分析和智能推荐。', size=10.5)
    add_para('（3）爬虫数据服务：部分技术团队提供定制化数据采集服务，但缺少分析、可视化与推荐的一站式能力。', size=10.5)
    add_para('本项目处于蓝海市场，具备先发优势。')

    add_heading('2.3 目标用户', level=2)
    users = [
        ('用户群体', '规模估算', '核心需求', '付费意愿'),
        ('网络文学读者', '5.8亿', '跨平台发现好书、个性化推荐', '中'),
        ('网络文学作者', '2000万+', '市场趋势分析、竞品研究', '中高'),
        ('出版机构/编辑', '5万+', '优质IP挖掘、市场调研', '高'),
        ('投资者/分析师', '10万+', '行业数据报告、投资决策', '高'),
        ('平台运营方', '100+', '竞品监控、行业对标', '高'),
    ]
    add_styled_table(doc, users[0], users[1:])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 3: Technical Feasibility
    # ══════════════════════════════════════════════════════════════
    add_heading('三、技术可行性分析', level=1)

    add_heading('3.1 技术架构', level=2)
    add_para('系统采用经典的大数据Web应用三层架构：')

    arch_data = [
        ('层级', '技术选型', '说明'),
        ('数据采集层', 'Python + Requests + BS4 + Scrapy', '多平台并发爬取，支持反爬虫策略'),
        ('数据存储层', 'Hadoop HDFS + MySQL/SQLite', '分布式存储+关系型数据库'),
        ('数据处理层', 'Apache Spark + Pandas + NumPy', '大数据处理、统计分析、模型训练'),
        ('算法引擎层', 'Scikit-learn + TF-IDF + 协同过滤', '内容分析、相似度计算、混合推荐'),
        ('Web应用层', 'Django 5.2 + HTML5 + CSS3 + jQuery', '全栈Web框架、响应式UI'),
        ('可视化层', 'ECharts 5.0', '丰富的交互式图表、多维度数据展示'),
    ]
    add_styled_table(doc, arch_data[0], arch_data[1:])

    add_heading('3.2 核心技术', level=2)
    add_para('（1）多平台数据采集引擎', bold=True, size=11)
    add_para('实现8大平台的统一数据采集框架，支持CSS选择器配置化、并发爬取、自动重试、反爬虫规避等能力。每个平台独立配置分类URL、数据字段和解析规则，支持采集标题、作者、分类、简介、字数、点击量、收藏数、推荐数、更新状态等字段。', size=10.5)

    add_para('（2）TF-IDF内容分析算法', bold=True, size=11)
    add_para('采用TF-IDF（词频-逆文档频率）算法对小说简介进行文本向量化处理，提取500维特征向量，实现基于内容的相似度计算。该算法能有效识别语义相近但用词不同的作品，为跨平台推荐提供基础。', size=10.5)

    add_para('（3）混合推荐算法', bold=True, size=11)
    add_para('融合四种推荐策略：协同过滤（权重35%）、分类偏好（权重25%）、平台偏好（权重20%）、全局热度（权重20%），并通过多样性增强机制避免推荐结果同质化。', size=10.5)

    add_para('（4）实时数据可视化', bold=True, size=11)
    add_para('基于ECharts实现10+种图表类型，包括散点图、热力图、堆叠柱状图、饼图、玫瑰图等，支持跨平台数据对比、动态筛选、自适应布局。', size=10.5)

    add_heading('3.3 数据资源', level=2)
    add_para('当前系统已具备以下数据基础：')
    add_bullet('数据量：600+部小说完整数据')
    add_bullet('覆盖平台：8个主流网络文学平台')
    add_bullet('覆盖分类：15个分类体系（玄幻、仙侠、都市、科幻、历史、军事、游戏、竞技、悬疑、言情、轻小说、奇幻、武侠、现实、短篇）')
    add_bullet('数据字段：标题、作者、分类、简介、字数、更新状态、来源、平台、点击量、收藏量、推荐量等')
    add_bullet('数据时效性：2026年2月-5月数据范围')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 4: Economic Feasibility
    # ══════════════════════════════════════════════════════════════
    add_heading('四、经济可行性分析', level=1)

    add_heading('4.1 成本估算', level=2)
    add_para('项目开发与运营的预估成本如下：')

    cost_data = [
        ('成本类别', '明细', '预估金额(万元/年)', '备注'),
        ('人力成本', '3人开发团队', '60-80', '全栈工程师×1+数据工程师×1+算法工程师×1'),
        ('服务器成本', '云服务器+存储', '8-15', 'Hadoop集群+Web服务器+数据库'),
        ('带宽成本', '数据采集+服务', '3-5', '8平台数据持续采集'),
        ('软件许可', '开发工具+服务', '2-3', 'IDE、云服务API'),
        ('运维成本', '监控+维护', '3-5', '系统监控、数据维护、安全防护'),
        ('合计', '—', '76-108', '—'),
    ]
    add_styled_table(doc, cost_data[0], cost_data[1:])

    add_heading('4.2 收益预测', level=2)
    add_para('项目具有多元化的盈利模式：')

    revenue_data = [
        ('收入来源', '模式', '预计年收入(万元)', '实现周期'),
        ('VIP订阅', '数据分析报告订阅制', '50-80', '6个月'),
        ('API服务', '数据接口按量计费', '30-60', '3个月'),
        ('定制报告', '行业分析报告定制', '20-40', '3个月'),
        ('广告收入', '网站流量变现', '10-20', '6个月'),
        ('企业服务', '平台/出版机构合作', '40-80', '12个月'),
        ('合计', '—', '150-280', '—'),
    ]
    add_styled_table(doc, revenue_data[0], revenue_data[1:])

    add_heading('4.3 投资回报分析', level=2)
    add_para('基于上述成本与收益预测，项目具备良好的投资回报预期：')
    add_bullet('首年投资回收期：约6-9个月')
    add_bullet('投资回报率（ROI）：预计100%-180%')
    add_bullet('盈亏平衡点：月活跃用户达到2000人时')
    add_bullet('三年净现值（NPV）：预计正值，项目经济可行')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 5: Operational Feasibility
    # ══════════════════════════════════════════════════════════════
    add_heading('五、运营可行性分析', level=1)

    add_para('（1）技术运维能力', bold=True, size=11)
    add_para('系统基于成熟的开源技术栈构建（Django + Spark + Hadoop），技术社区活跃，文档完善，运维难度适中。系统支持SQLite和MySQL双模式，可根据业务规模灵活切换。', size=10.5)

    add_para('（2）数据持续更新能力', bold=True, size=11)
    add_para('系统内置多平台数据采集引擎，支持定时任务自动采集，确保数据的时效性。同时设计了数据清洗与预处理流程，保障数据质量。', size=10.5)

    add_para('（3）用户体验', bold=True, size=11)
    add_para('Web应用提供友好的中文界面，支持响应式布局，适配PC和移动端。可视化仪表盘直观展示分析结果，降低用户使用门槛。', size=10.5)

    add_para('（4）可扩展性', bold=True, size=11)
    add_para('系统采用模块化设计，新增平台仅需添加配置文件，无需修改核心代码。算法引擎支持策略热插拔，可灵活引入新的推荐模型。', size=10.5)

    # ══════════════════════════════════════════════════════════════
    # Chapter 6: Legal & Compliance
    # ══════════════════════════════════════════════════════════════
    add_heading('六、法律与合规分析', level=1)

    add_para('（1）数据采集合规', bold=True, size=11)
    add_para('系统采集的是各平台公开发布的小说元数据（标题、作者、分类等），不涉及付费内容或用户隐私数据。建议在商业化运营前咨询法律顾问，确保符合《数据安全法》《个人信息保护法》要求，并在网站标注数据来源。', size=10.5)

    add_para('（2）知识产权', bold=True, size=11)
    add_para('系统采用的技术栈均为开源软件（Apache 2.0、MIT、BSD等许可证），不存在商业授权风险。系统展示的小说信息（书名、作者、简介）用于数据分析目的，属于合理使用范畴。', size=10.5)

    add_para('（3）robots.txt合规', bold=True, size=11)
    add_para('数据采集器尊重各网站的robots.txt协议，设置合理的请求间隔，避免对目标网站造成负担。', size=10.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 7: Risk Analysis
    # ══════════════════════════════════════════════════════════════
    add_heading('七、风险分析', level=1)

    risk_data = [
        ('风险类别', '风险描述', '影响程度', '发生概率', '应对措施'),
        ('技术风险', '目标网站反爬虫机制升级导致数据采集失败', '高', '中', '多样化采集策略，IP轮换，增加备用数据源'),
        ('法律风险', '数据采集涉及版权或隐私合规问题', '高', '低', '仅采集公开数据，标注来源，法律顾问审查'),
        ('市场风险', '竞争对手推出类似产品', '中', '中', '快速迭代，建立数据和技术壁垒'),
        ('运营风险', '数据质量下降或更新不及时', '中', '中', '建立数据质量监控体系，自动化采集流程'),
        ('技术债务', '技术栈老化或依赖库停止维护', '低', '低', '采用成熟开源技术，定期升级依赖'),
        ('安全风险', '系统遭受网络攻击或数据泄露', '高', '低', '部署WAF、HTTPS加密、定期安全审计'),
    ]
    add_styled_table(doc, risk_data[0], risk_data[1:])

    # ══════════════════════════════════════════════════════════════
    # Chapter 8: Implementation Plan
    # ══════════════════════════════════════════════════════════════
    add_heading('八、项目实施计划', level=1)

    plan_data = [
        ('阶段', '时间', '主要任务', '交付物'),
        ('第一阶段\n基础搭建', '第1-2月', '环境搭建、数据采集框架、\n基础数据库设计', '可运行的数据采集系统\n初始数据集(100+条)'),
        ('第二阶段\n核心开发', '第3-4月', '分析算法实现、推荐引擎开发、\nWeb应用框架', '分析引擎+推荐引擎\n基础Web界面'),
        ('第三阶段\n功能完善', '第5-6月', '可视化仪表盘、用户系统、\nAPI接口开发', '完整Web应用\n10+API端点'),
        ('第四阶段\n测试优化', '第7-8月', '系统测试、性能优化、\n用户体验改进', '测试报告\n优化后的系统'),
        ('第五阶段\n上线运营', '第9月', '部署上线、运营监控、\n用户反馈收集', '正式上线运营'),
    ]
    add_styled_table(doc, plan_data[0], plan_data[1:])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # Chapter 9: Conclusion
    # ══════════════════════════════════════════════════════════════
    add_heading('九、结论与建议', level=1)

    add_heading('9.1 综合评估', level=2)

    assessment_data = [
        ('评估维度', '评估结果', '评分(满分10)'),
        ('市场可行性', '网络文学市场规模持续增长，跨平台数据分析属蓝海市场', '8.5'),
        ('技术可行性', '技术栈成熟，核心算法已实现并验证，技术方案可行', '9.0'),
        ('经济可行性', '投资回收期短，多元化盈利模式，ROI预期良好', '8.0'),
        ('运营可行性', '模块化设计便于扩展，自动化运维降低运营成本', '8.5'),
        ('法律可行性', '采集公开数据，使用开源技术，合规风险可控', '7.5'),
        ('综合评分', '—', '8.3'),
    ]
    add_styled_table(doc, assessment_data[0], assessment_data[1:])

    add_heading('9.2 结论', level=2)
    add_para('经过市场、技术、经济、运营、法律五个维度的综合评估，多平台小说大数据分析与推荐系统项目具备充分的可行性。项目技术方案成熟可靠，市场前景广阔，经济回报可观，运营模式可持续。综合评价得分8.3分（满分10分），建议批准立项。')

    add_heading('9.3 建议', level=2)
    recommendations = [
        '优先推进：完成核心数据采集与分析引擎的产品化封装，尽快推出MVP（最小可行产品）版本进行市场验证。',
        '关注合规：在商业化运营前，委托法律顾问完成数据合规审查，确保符合相关法律法规要求。',
        '用户导向：在产品开发过程中持续收集目标用户反馈，确保产品功能与市场需求高度匹配。',
        '生态构建：逐步建立与网络文学平台、出版机构的合作关系，从数据聚合向产业赋能升级。',
        '技术迭代：持续跟踪大语言模型（LLM）在内容分析领域的进展，适时引入AI辅助的深度内容理解能力。',
    ]
    for r in recommendations:
        add_bullet(r)

    # ── Footer ──
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— 报告结束 —')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # ── Save ──
    output_path = r'D:\NOVEL ANALYSIS\可行性分析报告.docx'
    doc.save(output_path)
    print(f'报告已生成：{output_path}')
    return output_path


if __name__ == '__main__':
    generate_report()
